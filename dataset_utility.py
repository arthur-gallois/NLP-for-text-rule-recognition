import os
import json
import sys


class Text:
    '''
    Objet représentant un texte annoté sur word. 
    Il contient le texte complet ainsi qu'une liste des règles annotées (voir objet Rule). 
    Un print d'une instance de cette classe renvoie une structure en json
    '''

    rules = []
    text = ""

    _min_sentence_len = 10

    def __init__(self, text):
        self.text = text

    def __repr__(self):
        return '{' + f'"text" : "{self.text}","rules" : {self.rules}' + '}'

    def stringify(self):
        return self.__repr__()

    def get_sentences(self):
        '''
        Renvoie toutes les phrases de plus de _min_sentence_len du texte complet, classée selon la présence ou non d'une règle.
        '''
        sentences = [sentence.strip(" ") for sentence in self.text.split(
            '.') if len(sentence.strip(" ")) > self._min_sentence_len]
        conditions = [rule.text for rule in self.rules]
        non_conditional_sentences = []
        conditional_sentences = []
        for sentence in sentences:
            is_conditional = False
            n = 0
            for condition in conditions:
                if condition in sentence:
                    is_conditional = True
                    n += 1
            if is_conditional:
                conditional_sentences.append(sentence)
            else:
                non_conditional_sentences.append(sentence)
        return(conditional_sentences, non_conditional_sentences)


class Rule:
    '''
    Objet représentant une règle annotée.
    L'objet contient le texte complet de la règle ainsi que les parties condition, consequence et action
    '''
    text = ""
    condition = ""
    consequence = ""
    action = ""

    def __init__(self, text, condition, consequence, action):
        self.text = text
        self.condition = condition
        self.consequence = consequence
        self.action = action

    def __repr__(self):
        text = self.text
        condition = self.condition
        consequence = self.consequence
        action = self.action
        return '{' + f'"text" : "{text}", "condition" : "{condition}", "consequence" : "{consequence}","action" : "{action}"' + '}'


def _parse_json(input_json):
    '''
    Renvoie un objet Text à partir d'une string au format json le représentant
    '''
    try:
        text_dict = json.loads(input_json)
        text = Text(text_dict["text"])
        for rule in text_dict["rules"]:
            text.rules.append(Rule(rule["text"], rule["condition"],
                                   rule["consequence"], rule["action"]))
    except:
        print("\033[91m Error : Wrong text format \033[0m")
        return None
    return text


def getJson(json_name):
    '''
    Renvoie un objet Text à partir du nom d'un json (sans le .json) présent dans le dossier ./dataset/json/
    '''
    file_path = os.path.dirname(os.path.abspath(__file__))
    json_path = f"{file_path}/dataset/json/{json_name}.json"
    f = open(json_path, "r", encoding='utf-8')
    lines = f.readlines()
    return _parse_json(lines[0])


def _word_loader(documentWord):
    '''
    Renvoie un objet Text à partir du chemin d'accès du word
    '''
    import docx

    def _clear_text(*texts):
        '''
        Envoie d'une string tout caractère pouvant gêner le .json
        '''
        L = []
        for text in texts:
            text = text.replace('\\', '')
            text = text.replace('"', '')
            text = text.replace("‘", '')
            text = text.replace("’", '')
            text = text.replace("'", '')
            text = text.replace('\n', '')
            text = text.replace("\xa0", '')
            L.append(text)

        if(len(L) == 1):
            return(L[0])
        return L
    # documentWord est le chemin du docx (au format str)
    doc = docx.Document(documentWord)
    k = len(doc.paragraphs)
    word_text = ""
    for para in doc.paragraphs: # Parcours les paragraphes du docx pour en extraire le texte complet
        text = para.text
        text = _clear_text(text)
        word_text += text
    text = Text(word_text)

    rules = []
    for para in doc.paragraphs:
        rule = ""
        condition = ""
        consequence = ""
        action = ""
        for run in para.runs:
            if run.underline:
                rule += run.text
            if run.bold and not run.italic:
                condition += run.text
            if run.italic and not run.bold:
                consequence += run.text
            if run.bold and run.italic:
                action += run.text
            if not run.underline and rule != "":
                rule, condition, consequence, action = _clear_text(
                    rule, condition, consequence, action)
                rules.append(
                    Rule(rule.strip("."), condition, consequence, action))
                rule = ""
                condition = ""
                consequence = ""
                action = ""
        if rule != "" and condition != "":
            rule, condition, consequence, action = _clear_text(
                rule, condition, consequence, action)
            rules.append(Rule(rule.strip(".").replace('\n', '').replace(
                "\xa0", ''), condition, consequence, action))
            rule = ""
            condition = ""
            consequence = ""
            action = ""
    text.rules = rules
    return text


def _save_json(json_name, json_text):
    '''
    Fonction de sauvegarde de fichier texte. Ajoute un .json à json_name
    '''
    path = os.path.dirname(os.path.abspath(__file__))
    with open(f"{path}/dataset/json/{json_name}.json", 'w', encoding='utf8') as file:
        file.write(json_text)


if __name__ == '__main__':
    if(len(sys.argv) < 2):
        print("\033[91m Veuillez spécifier le nom du word \033[0m")
    elif(sys.argv[1] == "all"): # Si l'argument est all, regénère les json de tous les docx présents dans ./dataset/docx
        path = os.path.dirname(os.path.abspath(__file__))
        for _, _, filenames in os.walk(f'{path}/dataset/docx/'):
            for filename in filenames:
                name = filename.split(".")[0]
                print(f"Generating {name}.json")
                json_text = _word_loader(f"{path}/dataset/docx/{name}.docx")
                _save_json(name, json_text.stringify())
                test = _parse_json(json_text.stringify())
    else: #Sinon génère uniquement le docx du fichier voulu
        path = os.path.dirname(os.path.abspath(__file__))
        json_text = _word_loader(f"{path}/dataset/docx/{sys.argv[1]}.docx")
        _save_json(sys.argv[1], json_text.stringify())
        test = _parse_json(json_text.stringify())
