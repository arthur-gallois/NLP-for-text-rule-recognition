import docx


def wordLoader(documentWord):

    doc = docx.Document(documentWord)

    k = len(doc.paragraphs)

    bolds = [[] for i in range(k)]

    italics = [[] for i in range(k)]

    p = 0

    for para in doc.paragraphs:

        for run in para.runs:

            if run.italic:

                italics[p].append(run.text)

            if run.bold:

                bolds[p].append(run.text)

        p += 1

    print(bolds)


doc = docx.Document("test_doc.docx")
k = len(doc.paragraphs)
bolds = [[] for i in range(k)]
italics = [[] for i in range(k)]
p = 0
for para in doc.paragraphs:
    for run in para.runs:
        if run.italic:
            italics[p].append(run.text)
        if run.bold:
            bolds[p].append(run.text)
    p += 1
print(bolds)
